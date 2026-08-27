import os
import sys
import json
import time
import re
import shutil
import signal
import subprocess
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from enum import Enum

import numpy as np
import pandas as pd

from src.paper.tex_exporter import MarkdownToTexConverter
from src.agents.crew import Agent, Task, Crew, Process

# =============================================================================
# 多 LLM 提供商支持
# =============================================================================

try:
    from src.llm import get_provider_manager, ProviderType, LLMProviderFactory, ProviderConfig
    _LLM_PROVIDER_AVAILABLE = True
except ImportError:
    _LLM_PROVIDER_AVAILABLE = False


_llm_provider_instance = None


def _load_claude_settings_env() -> None:
    """将 Claude Code 的 settings.json 中的 env 注入 os.environ。

    覆盖顺序（低→高）：
    1. 全局 ~/.claude/settings.json
    2. 项目级 .claude/settings.json
    3. 项目级 .claude/settings.local.json
    已存在的环境变量不覆盖，方便用户显式指定。
    """
    if os.environ.get("DEFAULT_LLM_PROVIDER", "").lower() == "claude_cli":
        return  # 用户显式要求 CLI，则不注入
    candidates = [
        Path.home() / ".claude" / "settings.json",
        Path.cwd() / ".claude" / "settings.json",
        Path.cwd() / ".claude" / "settings.local.json",
    ]
    for path in candidates:
        try:
            if not path.exists():
                continue
            data = json.loads(path.read_text(encoding="utf-8"))
            env = data.get("env") or {}
            for k, v in env.items():
                if isinstance(v, str) and k not in os.environ:
                    os.environ.setdefault(k, v)
        except (OSError, ValueError, TypeError):
            continue


def _get_llm_provider():
    """获取 LLM Provider 实例（惰性初始化）

    默认优先使用 API Provider；只有显式配置 DEFAULT_LLM_PROVIDER=claude_cli
    或完全没有任何可用 API 配置时才回退到 Claude CLI。
    """
    global _llm_provider_instance
    if _llm_provider_instance is not None:
        return _llm_provider_instance

    if not _LLM_PROVIDER_AVAILABLE:
        return None

    # 先注入 .claude/settings.json 的 env，使 API Provider 探测可用
    _load_claude_settings_env()

    if not _LLM_PROVIDER_AVAILABLE:
        return None

    try:
        default_provider = os.getenv("DEFAULT_LLM_PROVIDER", "").lower()

        if default_provider == "claude_cli":
            manager = get_provider_manager()
            manager.register(ProviderType.CLAUDE_CLI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Claude CLI Provider (model={manager.get().config.model})")
        elif default_provider == "openai" or os.getenv("OPENAI_API_KEY"):
            manager = get_provider_manager()
            manager.register(ProviderType.OPENAI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 OpenAI Provider (model={manager.get().config.model})")
        elif default_provider == "anthropic" or os.getenv("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_AUTH_TOKEN"):
            manager = get_provider_manager()
            config = ProviderConfig.from_env(ProviderType.ANTHROPIC)
            if not config.api_key and os.getenv("ANTHROPIC_AUTH_TOKEN"):
                config.api_key = os.getenv("ANTHROPIC_AUTH_TOKEN")
            if os.getenv("ANTHROPIC_BASE_URL"):
                config.api_host = os.getenv("ANTHROPIC_BASE_URL")
                # 代理模式下设置5分钟超时
                config.timeout = min(config.timeout, 300)
            if os.getenv("ANTHROPIC_MODEL"):
                config.model = os.getenv("ANTHROPIC_MODEL")
            manager.register(ProviderType.ANTHROPIC, config)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Anthropic Provider (model={manager.get().config.model})")
        elif default_provider == "gemini" or os.getenv("GEMINI_API_KEY"):
            manager = get_provider_manager()
            manager.register(ProviderType.GEMINI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Gemini Provider (model={manager.get().config.model})")
        elif default_provider == "ollama" or os.getenv("OLLAMA_MODEL") or os.getenv("OLLAMA_HOST"):
            manager = get_provider_manager()
            manager.register(ProviderType.OLLAMA)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Ollama Provider (model={manager.get().config.model})")
        elif default_provider == "claude_cli" or (
            not os.getenv("OPENAI_API_KEY")
            and not os.getenv("ANTHROPIC_AUTH_TOKEN")
            and not os.getenv("ANTHROPIC_API_KEY")
            and not os.getenv("GEMINI_API_KEY")
        ):
            # 显式要求 CLI，或完全无可用的 API 配置时才走 Claude CLI
            manager = get_provider_manager()
            manager.register(ProviderType.CLAUDE_CLI)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Claude CLI Provider (model={manager.get().config.model})")
        else:
            # 默认优先走 Anthropic 兼容 API（如 litellm 代理）
            manager = get_provider_manager()
            config = ProviderConfig.from_env(ProviderType.ANTHROPIC)
            if not config.api_key and os.getenv("ANTHROPIC_AUTH_TOKEN"):
                config.api_key = os.getenv("ANTHROPIC_AUTH_TOKEN")
            if os.getenv("ANTHROPIC_BASE_URL"):
                config.api_host = os.getenv("ANTHROPIC_BASE_URL")
                config.timeout = min(config.timeout, 300)
            if os.getenv("ANTHROPIC_MODEL"):
                config.model = os.getenv("ANTHROPIC_MODEL")
            manager.register(ProviderType.ANTHROPIC, config)
            _llm_provider_instance = manager
            print(f"[LLM] 使用 Anthropic API Provider (model={manager.get().config.model})")
    except Exception as e:
        print(f"[LLM] Provider 初始化失败: {e}，将回退到 Claude CLI")
        _llm_provider_instance = None

    return _llm_provider_instance


def _find_claude_code() -> Optional[str]:
    """自动搜索 Claude Code CLI 路径"""
    found = shutil.which("claude-code") or shutil.which("claude")
    return found


def _call_llm(
    prompt: str,
    system_prompt: Optional[str] = None,
    model: str = "sonnet",
    timeout: int = 600,
    max_retries: int = 3,
    retry_wait: int = 60
) -> str:
    """
    调用 LLM 生成内容
    优先使用配置的 API Provider，回退到 Claude Code CLI
    """
    provider_manager = _get_llm_provider()
    api_timed_out = False  # 标记 API 是否已超时，用于快速回退

    if provider_manager is not None:
        for attempt in range(max_retries):
            # API 首次超时后立即回退，避免用户长时间等待
            if api_timed_out:
                print("    [LLM] API Provider 不可用，直接回退到 Claude CLI...")
                break

            try:
                if attempt > 0:
                    print(f"    [LLM调用 {attempt + 1}/{max_retries}] 等待{retry_wait}秒后重试...")
                    time.sleep(retry_wait)

                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 开始请求...")

                # 使用 alarm 机制防止 API 调用卡死（仅 Linux）
                result = []
                exception = []
                def _handler(signum, frame):
                    raise TimeoutError("LLM API 调用超时")

                old_handler = None
                use_alarm = hasattr(signal, 'SIGALRM')
                if use_alarm:
                    old_handler = signal.signal(signal.SIGALRM, _handler)
                    signal.alarm(timeout)

                try:
                    response = provider_manager.generate(
                        prompt=prompt,
                        system_prompt=system_prompt,
                        timeout=timeout
                    )
                    result.append(response)
                except TimeoutError:
                    api_timed_out = True
                    raise
                finally:
                    if use_alarm:
                        signal.alarm(0)
                        if old_handler is not None:
                            signal.signal(signal.SIGALRM, old_handler)

                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 成功!")
                return result[0]
            except (TimeoutError, TimeoutError):
                api_timed_out = True
                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 超时，回退到 Claude CLI")
                break
            except Exception as e:
                print(f"    [LLM调用 {attempt + 1}/{max_retries}] 失败: {str(e)[:200]}")
                if attempt < max_retries - 1:
                    continue
                print("    [LLM] API Provider 失败，回退到 Claude CLI...")
                break

    # 回退到 Claude Code CLI
    claude_path = _find_claude_code()
    if not claude_path:
        raise RuntimeError(
            "Claude Code CLI 未找到，请确保已安装 Claude Code 并添加到 PATH，"
            "或配置 OPENAI_API_KEY / ANTHROPIC_API_KEY / GEMINI_API_KEY 环境变量"
        )

    full_prompt = prompt
    if system_prompt:
        full_prompt = f"{system_prompt}\n\n{prompt}"

    cmd = [
        claude_path,
        "-p",
        "--model", model,
        "--output-format", "json",
        full_prompt
    ]

    for attempt in range(max_retries):
        try:
            if attempt > 0:
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 等待{retry_wait}秒后重试...")
                time.sleep(retry_wait)

            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 开始请求...")
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            stdout, stderr = proc.communicate(timeout=timeout)

            stdout_text = stdout.decode("utf-8", errors="replace").strip()

            if proc.returncode != 0:
                error_msg = stderr.decode("utf-8", errors="replace")
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 失败: {error_msg[:200]}")
                if attempt < max_retries - 1:
                    continue
                raise RuntimeError(f"LLM 调用失败: {error_msg[:500]}")

            try:
                data = json.loads(stdout_text.strip())
            except json.JSONDecodeError:
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
                return stdout_text.strip()

            result_text = data.get("result", "")
            if isinstance(result_text, str):
                result_text = result_text.strip()
                if result_text.startswith("```"):
                    lines = result_text.splitlines()
                    if lines:
                        if lines[0].startswith("```"):
                            lines = lines[1:]
                        if lines and lines[-1].strip() == "```":
                            lines = lines[:-1]
                    result_text = "\n".join(lines).strip()
                print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
                return result_text
            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 成功!")
            return str(result_text)

        except subprocess.TimeoutExpired:
            print(f"    [LLM调用(CLI) {attempt + 1}/{max_retries}] 超时({timeout}秒)")
            if attempt < max_retries - 1:
                continue
            raise RuntimeError(f"LLM 调用超时（{timeout}秒），已重试{max_retries}次")
        except FileNotFoundError:
            raise RuntimeError("Claude Code CLI 未找到")

    raise RuntimeError("LLM 调用失败，已达到最大重试次数")


# =============================================================================
# 统一工作流引擎
# =============================================================================

from src.workflow import (
    Coordinator,
    DependencyType,
    MathModelingTemplate,
    CourseworkTemplate,
    FinancialAnalysisTemplate,
    get_template,
    CritiqueEngine,
    CodeExecutor,
    PaperGenerator,
)
from src.document_processing import DocumentLoader


try:
    from src.knowledge import KnowledgeBase
    _KNOWLEDGE_AVAILABLE = True
except ImportError:
    _KNOWLEDGE_AVAILABLE = False

try:
    from src.knowledge.algorithm_library import get_algorithm_library
    _ALGO_LIBRARY_AVAILABLE = True
except ImportError:
    _ALGO_LIBRARY_AVAILABLE = False


class UnifiedWorkflow:
    """
    统一工作流引擎 v2.0

    四阶段架构（借鉴 LLM-MM-Agent）：
    Stage 1: Problem Analysis - 问题分析 + DAG构建
    Stage 2: Mathematical Modeling - 数学建模（按DAG拓扑序）
    Stage 3: Computational Solving - 计算求解（Claude CLI代码）
    Stage 4: Paper Generation - 论文生成（大纲驱动 + 相关性过滤）
    """

    def __init__(
        self,
        output_dir: str = "work",
        template_name: str = "math_modeling",
        use_knowledge_base: bool = True,
        use_critique: bool = True,
    ):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.template_name = template_name
        self.use_critique = use_critique
        self.use_knowledge_base = use_knowledge_base and _KNOWLEDGE_AVAILABLE

        # 初始化各组件
        self.coordinator = Coordinator()
        self.critique_engine = CritiqueEngine(_call_llm)
        self.code_executor = CodeExecutor(
            call_llm=_call_llm,
            output_dir=str(self.output_dir),
        )
        self.paper_generator = PaperGenerator(
            call_llm=_call_llm,
            template=get_template(template_name),
            output_dir=str(self.output_dir),
        )
        self.document_loader = DocumentLoader()
        self.knowledge_base = KnowledgeBase() if self.use_knowledge_base else None

        # ======================================================================
        # 算法知识库：集成 Algorithms_MathModels 仓库的算法
        # ======================================================================
        self.algorithm_library = None
        if _ALGO_LIBRARY_AVAILABLE:
            try:
                self.algorithm_library = get_algorithm_library()
                print(f"[AlgorithmLibrary] 算法知识库已加载 ({len(self.algorithm_library.categories)} 个类别)")
            except Exception as e:
                print(f"[AlgorithmLibrary] 加载失败: {e}")

        # 全局上下文
        self.context: Dict[str, Any] = {}
        self.problem_text = ""
        self.data_files: Dict[str, str] = {}

        # Claude CLI path for code/algorithm generation
        self._claude_path = _find_claude_code()

        # ======================================================================
        # 显式记忆池：每个阶段完成后生成结构化摘要，供下一阶段调用
        # ======================================================================
        self.memory_pool: Dict[str, str] = {
            "analysis_summary": "",      # 阶段1摘要
            "modeling_summary": "",      # 阶段2摘要
            "algorithm_summary": "",     # 阶段3算法摘要
            "results_summary": "",       # 阶段3结果摘要
            "validation_summary": "",    # 阶段3b 回测/OOS 验证摘要（CCF-A）
            "reproducibility_summary": "",# 阶段3c 复现性打包摘要（CCF-A）
            "chapter_summaries": {},     # 论文各章摘要，key=chapter_id
        }

        # Per-agent persistent memory system
        try:
            from src.memory import MemorySystem
            self.memory_system = MemorySystem(str(self.output_dir))
            print(f"[MemorySystem] 已初始化，目录: {self.output_dir}/memory/")
        except ImportError:
            self.memory_system = None
            print("[MemorySystem] 模块不可用，将使用内存记忆池")

    def run_full_workflow(
        self,
        problem_text: str,
        data_files: Dict[str, str],
        problem_name: str = "数学建模问题",
    ) -> str:
        """运行完整工作流，输出 Markdown + Word + LaTeX 三格式论文。"""
        print("\n" + "=" * 70)
        print("数学建模论文自动生成系统 v2.3")
        print("架构: CrewAI Agent 协作 + 显式记忆池")
        print(f"模板: {self.template_name}")
        print("=" * 70)

        self.problem_text = problem_text
        self.data_files = data_files
        self.context["problem_text"] = problem_text
        self.context["data_files"] = data_files

        # Create role-based agents (crewAI style) with persistent memory
        llm_cb = lambda p, s: _call_llm(p, s)
        ms = self.memory_system
        analyst = Agent(role="问题分析师", goal="深入分析赛题，提取子任务与约束", llm_callback=llm_cb, memory_system=ms)
        modeler = Agent(role="数学建模师", goal="建立严谨的数学模型与公式体系", llm_callback=llm_cb, memory_system=ms)
        solver = Agent(role="求解工程师", goal="设计算法并生成可执行代码", llm_callback=llm_cb, memory_system=ms)
        writer = Agent(role="论文写作专家", goal="撰写完整的数学建模竞赛论文", llm_callback=llm_cb, memory_system=ms)
        manager = Agent(role="协调者", goal="统筹各阶段工作，确保质量与衔接", allow_delegation=True, llm_callback=llm_cb, memory_system=ms)

        shared: Dict[str, Any] = {}

        def do_analysis(_task: str, _ctx: str) -> str:
            analysis = self._stage_problem_analysis()
            self.context["analysis"] = analysis
            self.memory_pool["analysis_summary"] = self._summarize_analysis(analysis)
            self._save_text("stage_1_analysis/analysis_summary.md", self.memory_pool["analysis_summary"])
            if ms:
                ms.save_agent_summary("问题分析师", self.memory_pool["analysis_summary"])
                ms.save_shared_summary("analysis_summary.md", self.memory_pool["analysis_summary"])
                ms.store_shared("analysis_summary", self.memory_pool["analysis_summary"], metadata={"stage": "1"})
            shared["analysis"] = analysis
            shared["analysis_summary"] = self.memory_pool["analysis_summary"]
            return self.memory_pool["analysis_summary"]

        def do_modeling(_task: str, _ctx: str) -> str:
            analysis = shared.get("analysis", self.context.get("analysis", {}))
            modeling = self._stage_mathematical_modeling(analysis)
            self.context["modeling"] = modeling
            self.memory_pool["modeling_summary"] = self._summarize_modeling(modeling)
            self._save_text("stage_2_modeling/modeling_summary.md", self.memory_pool["modeling_summary"])
            if ms:
                ms.save_agent_summary("数学建模师", self.memory_pool["modeling_summary"])
                ms.save_shared_summary("modeling_summary.md", self.memory_pool["modeling_summary"])
                ms.store_shared("modeling_summary", self.memory_pool["modeling_summary"], metadata={"stage": "2"})
            shared["modeling"] = modeling
            shared["modeling_summary"] = self.memory_pool["modeling_summary"]
            return self.memory_pool["modeling_summary"]

        def do_solving(_task: str, _ctx: str) -> str:
            modeling = shared.get("modeling", self.context.get("modeling", {}))
            solving = self._stage_computational_solving(modeling)
            self.context["execution_result"] = solving.get("execution_result", {})
            self.context["code"] = solving.get("code", "")
            self.context["result_analysis"] = solving.get("interpretation", "")
            self.memory_pool["algorithm_summary"] = self._summarize_algorithm(solving)
            self.memory_pool["results_summary"] = self._summarize_results(solving)
            self._save_text("stage_3_algorithm/algorithm_summary.md", self.memory_pool["algorithm_summary"])
            self._save_text("stage_6_result_analysis/results_summary.md", self.memory_pool["results_summary"])

            # CCF-A: Stage 3b 样本外回测 + Stage 3c 复现性打包
            if self._is_ccf_a_template():
                print("  [CCF-A] 启动 Stage 3b 回测 + Stage 3c 复现性打包...")
                validation = self._stage_validation_backtest(solving)
                reproducibility = self._stage_reproducibility_packaging(solving, validation)
                self.context["validation"] = validation
                self.context["reproducibility"] = reproducibility
                shared["validation"] = validation
                shared["reproducibility"] = reproducibility

            if ms:
                ms.save_agent_summary("求解工程师", self.memory_pool["algorithm_summary"])
                ms.save_shared_summary("results_summary.md", self.memory_pool["results_summary"])
                ms.store_shared("results_summary", self.memory_pool["results_summary"], metadata={"stage": "3"})
            shared["solving"] = solving
            shared["algorithm_summary"] = self.memory_pool["algorithm_summary"]
            shared["results_summary"] = self.memory_pool["results_summary"]
            return self.memory_pool["results_summary"]

        def do_paper(_task: str, _ctx: str) -> str:
            paper = self._stage_paper_generation_v2()
            if ms:
                ms.save_agent_summary("论文写作专家", paper[:2000])
                ms.save_shared_summary("paper_summary.md", paper[:2000])
            shared["paper"] = paper
            return paper

        analyst.execute = do_analysis
        modeler.execute = do_modeling
        solver.execute = do_solving
        writer.execute = do_paper

        tasks = [
            Task(description="分析问题", agent=analyst, output_key="analysis_summary"),
            Task(description="建立数学模型", agent=modeler, output_key="modeling_summary"),
            Task(description="计算求解", agent=solver, output_key="results_summary"),
            Task(description="撰写论文", agent=writer, output_key="paper"),
        ]

        process_mode = os.getenv("CREW_PROCESS_MODE", "sequential").lower()
        process_map = {
            "sequential": Process.SEQUENTIAL,
            "hierarchical": Process.HIERARCHICAL,
            "consensus": Process.CONSENSUS,
        }
        selected_process = process_map.get(process_mode, Process.HIERARCHICAL)
        print(f"[Crew] 流程模式: {selected_process.value}")

        crew = Crew(
            agents=[analyst, modeler, solver, writer, manager],
            tasks=tasks,
            process=selected_process,
            manager_agent=manager if selected_process == Process.HIERARCHICAL else None,
            shared_memory=shared,
        )
        crew.kickoff()

        paper = shared.get("paper", "")
        self._save_text("stage_6_result_analysis/results_summary.md", self.memory_pool.get("results_summary", ""))

        # 保存结果
        paper_path = self.output_dir / "final" / "MathModeling_Paper.md"
        paper_path.parent.mkdir(parents=True, exist_ok=True)
        paper_path.write_text(paper, encoding="utf-8")

        # 字数统计
        chinese_chars = len(re.findall(r'[一-鿿]', paper))
        print(f"\n{'='*60}")
        print("论文生成完成")
        print(f"{'='*60}")
        print(f"论文文件: {paper_path}")
        print(f"总字符数: {len(paper)}")
        print(f"中文字数: {chinese_chars}")

        # 导出解决方案
        solution_path = self.output_dir / "final" / "solution.json"
        self.coordinator.export_solution(solution_path)

        # 导出 docx
        self._convert_to_docx(paper_path)

        # 导出 LaTeX
        self._convert_to_tex(paper_path)

        self._save_json("final/memory_pool.json", self.memory_pool)
        if self.memory_system:
            self.memory_system.save_shared_json("memory_pool.json", self.memory_pool)

        return paper

    # ========================================================================
    # CCF-A 模板判定与严谨性清单
    # ========================================================================

    def _is_ccf_a_template(self) -> bool:
        """当前模板是否为 CCF-A 顶会模板（决定是否触发 rigor 批判/回测/打包）。"""
        try:
            return bool(getattr(self.paper_generator.template, "is_ccf_a", False))
        except Exception:
            return False

    def _get_rigor_checklist(self) -> list:
        """获取当前模板的 CCF-A 严谨性清单（非 CCF-A 返回空）。"""
        try:
            return list(self.paper_generator.template.get_rigor_checklist() or [])
        except Exception:
            return []

    # ========================================================================
    # Stage 1: 问题分析
    # ========================================================================

    def _stage_problem_analysis(self) -> Dict[str, Any]:
        """
        问题分析阶段

        借鉴 LLM-MM-Agent 的反思链设计 + Critique-Improvement
        + AI-Scientist v2 风格的创意性研究视角生成
        """
        print("  分析问题并构建DAG...")

        # 加载数据文件摘要
        data_descriptions = self._load_data_descriptions()
        self.context["data_descriptions"] = data_descriptions

        # 注册知识库（如果有相关文档）
        if self.knowledge_base and data_descriptions:
            for name, desc in data_descriptions.items():
                self.knowledge_base.add_document(
                    title=f"数据文件: {name}",
                    content=desc,
                    metadata={"type": "data"},
                )

        # ======================================================================
        # 创意性研究视角生成（AI-Scientist v2 风格 ideation pre-step）
        # ======================================================================
        print("  [Ideation] 生成创意研究视角...")
        try:
            from src.workflow.ideation import ProblemIdeation
            ideation = ProblemIdeation(_call_llm)
            data_desc_text = "\n".join(data_descriptions.values())[:1000]
            ideas = ideation.generate_ideas(
                self.problem_text,
                data_descriptions=data_desc_text,
                num_ideas=5,
                force_structural=self._is_ccf_a_template(),
            )
            ideation_context = ideation.format_for_analysis(ideas)
            self.context["ideation_ideas"] = ideas
            self.context["ideation_context"] = ideation_context
            if self.memory_system:
                self.memory_system.save_shared_json("ideation_results.json", ideas)
                self.memory_system.store_shared(
                    "ideation_ideas",
                    ideation_context,
                    metadata={"stage": "ideation", "count": len(ideas)},
                )
            print(f"  [Ideation] 已生成 {len(ideas)} 个研究视角")
        except Exception as e:
            print(f"  [Ideation] 生成失败: {e}，将跳过创意视角注入")
            ideation_context = ""
            self.context["ideation_ideas"] = []

        # Prompt: 反思链设计（借鉴 LLM-MM-Agent PROBLEM_ANALYSIS_PROMPT）
        ideation_section = f"\n【研究视角建议 - 供参考的创意方向】\n{ideation_context}\n" if ideation_context else ""

        prompt = f"""请对以下数学建模问题进行深度分析。

【赛题】
{self.problem_text[:4000]}

【数据文件描述】
{chr(10).join(data_descriptions.values())[:2000]}
{ideation_section}
要求：
1. 识别问题的核心组件和它们之间的依赖关系
2. 分析问题的动态特性（时间演化、空间分布等）
3. 从多个视角审视问题（数学视角、物理视角、工程视角）
4. 识别关键假设和潜在的不确定性
5. 将问题分解为可管理的子任务

输出严格的JSON格式：
{{
  "background": "问题背景概述（300字）",
  "sub_problems": [
    {{
      "id": "task_1",
      "description": "子任务描述",
      "objective": "任务目标",
      "key_constraints": ["约束1", "约束2"],
      "suggested_methods": ["建议方法1", "建议方法2"]
    }}
  ],
  "key_assumptions": ["假设1", "假设2"],
  "data_dependencies": {{"task_1": ["数据文件1"]}},
  "solution_approach": "整体解决思路（500字）"
}}"""

        # 生成 + Critique-Improvement
        print("  生成初始分析...")
        result = _call_llm(prompt, "你是一位数学建模专家，擅长深度问题分析。")

        analysis = self._parse_json_safely(result, self._default_analysis())

        if self.use_critique:
            print("  启动Critique-Improvement循环...")
            analysis_text = json.dumps(analysis, ensure_ascii=False, indent=2)
            improved = self.critique_engine.critique_and_improve(
                content=analysis_text,
                content_type="analysis",
                context=self.problem_text[:2000],
                max_iterations=1,
                score_threshold=8.0,
            )
            analysis = self._parse_json_safely(improved, analysis)

        # 构建DAG
        sub_problems = analysis.get("sub_problems", [])
        for i, sp in enumerate(sub_problems):
            task_id = sp.get("id", f"task_{i+1}")
            deps = {}
            # 简单的顺序依赖：后面的任务依赖前面的
            if i > 0:
                prev_id = sub_problems[i-1].get("id", f"task_{i}")
                deps[prev_id] = [DependencyType.STRUCTURAL, DependencyType.DATA]
            self.coordinator.register_task(
                task_id=task_id,
                description=sp.get("description", ""),
                dependencies=deps,
            )

        self.coordinator.analyze_dependencies()
        self.context["sub_problems"] = sub_problems

        # 保存
        self._save_json("stage_1_analysis/analysis.json", analysis)
        return analysis

    def _load_data_descriptions(self) -> Dict[str, str]:
        """加载数据文件描述"""
        descriptions = {}
        for name, filepath in self.data_files.items():
            path = Path(filepath)
            if not path.exists():
                descriptions[name] = f"{name}: 文件不存在"
                continue

            try:
                result = self.document_loader.load(filepath)
                if result.success and result.document:
                    desc = result.document.text[:1500]
                    descriptions[name] = f"【{name}】\n{desc}\n"
                    # 添加到知识库
                    if self.knowledge_base:
                        self.knowledge_base.add_document(
                            title=f"数据: {name}",
                            content=desc,
                            source=filepath,
                            metadata={"type": "data_file", "filename": name},
                        )
                else:
                    descriptions[name] = f"{name}: 加载失败 - {result.error}"
            except Exception as e:
                descriptions[name] = f"{name}: 错误 - {str(e)}"

        return descriptions

    # ========================================================================
    # Stage 2: 数学建模
    # ========================================================================

    def _stage_mathematical_modeling(self, analysis: Dict) -> Dict[str, Any]:
        """
        数学建模阶段 - 按DAG拓扑序执行每个子任务

        借鉴 LLM-MM-Agent：
        - 依赖上下文自动拼接
        - 方法检索增强生成
        - Actor-Critic-Improvement 循环
        """
        sub_problems = analysis.get("sub_problems", [])
        if not sub_problems:
            sub_problems = [{"id": "task_1", "description": "建立模型求解"}]

        # 处理所有子任务的详细建模（确保不遗留pending任务）
        tasks_to_model = self.coordinator.dag_order
        print(f"  共 {len(tasks_to_model)} 个子任务，将全部进行详细建模")

        all_formulas = []
        all_models = []

        for task_id in tasks_to_model:
            task_node = self.coordinator.tasks.get(task_id)
            if not task_node:
                continue

            print(f"\n  建模任务: {task_id} - {task_node.description[:60]}...")

            # 获取依赖上下文
            dep_context = self.coordinator.get_dependency_context(
                task_id, max_chars=3000
            )

            # 知识库检索相关方法
            kb_context = ""
            if self.knowledge_base:
                kb_results = self.knowledge_base.query_with_context(
                    task_node.description, top_k=3, max_chars=1500
                )
                if kb_results:
                    kb_context = f"【相关知识】\n{kb_results}\n\n"

            # ======================================================================
            # 算法知识库检索：自动推荐适用的数学建模算法
            # ======================================================================
            algo_context = ""
            if self.algorithm_library:
                print(f"    [AlgorithmLibrary] 检索相关算法...")
                algo_recommendation = self.algorithm_library.generate_recommendation_text(
                    query=task_node.description + " " + self.problem_text[:500],
                    top_k=3
                )
                if algo_recommendation:
                    algo_context = f"【算法库推荐】\n{algo_recommendation}\n\n"
                    print(f"    [AlgorithmLibrary] 已推荐 {len(self.algorithm_library.search(task_node.description + ' ' + self.problem_text[:500], top_k=3))} 个相关算法")

            # 1. 任务分析（缩短prompt以加快响应）
            analysis_prompt = f"""对以下子任务进行数学分析：

任务: {task_node.description}

背景: {self.problem_text[:1500]}

{dep_context}

{kb_context}

{algo_context}

分析核心数学结构、建模方法、关键变量。输出至少600字。"""

            task_analysis = _call_llm(
                analysis_prompt,
                "你是数学建模专家，擅长将实际问题抽象为数学结构。"
            )

            # 2. 公式生成（Actor）
            # CCF-A 模板：强制要求网络/反馈/异质性/因果识别结构，避免线性标定
            ccf_a_block = ""
            if self._is_ccf_a_template():
                ccf_a_block = """
【CCF-A 严谨性硬性要求 — 必须在模型中体现，缺失即拒稿】
1. 因果识别：明确区分"已识别"与"假设"的弹性；处理内生性/反向因果（如信贷供给↔房价）
2. 异质性：使用层级/面板结构刻画跨单元（省份/银行/企业）差异，禁止仅全国聚合
3. 反馈环：编码内生反馈（如银行资本约束→信贷供给→住房需求→房价→地价→城投→银行）
4. 网络传染：若涉及系统性风险，必须包含网络损失传染层（DebtRank / SDECM / 扩散谱 PDE / TGNN），并与聚合块对比
5. 形式化：所有构造量（如"土地净收入"、政策乘数）必须给出从原始输入到使用值的对账公式
6. 参数先验：列出所有参数的先验分布、取值范围、来源（用于后续复现性打包）
"""

            formulas_prompt = f"""基于以下分析建立数学模型。

任务分析:
{task_analysis[:1200]}

{algo_context}
{ccf_a_block}
要求：
1. 定义变量和参数（表格形式）
2. 建立核心数学公式（LaTeX格式，公式编号）
3. 说明公式物理意义
4. 列出模型假设
5. 若算法库推荐了适用的算法，请在模型中明确采用并说明理由

输出至少1200字的完整建模过程。"""

            formulas = _call_llm(
                formulas_prompt,
                "你是数学建模专家，擅长严谨的形式化建模。"
            )

            # 3. Critique-Improvement
            # CCF-A 模板：用 ccf_a_rigor 维度批判建模本身（而非仅 paper_chapter），
            # 把"模型建错了"在 Stage 2 就拦截下来，避免论文阶段才暴露。
            if self.use_critique:
                if self._is_ccf_a_template():
                    print(f"    [CCF-A] 对建模进行 ccf_a_rigor 批判...")
                    formulas = self.critique_engine.critique_and_improve(
                        content=formulas,
                        content_type="ccf_a_rigor",
                        context=f"任务: {task_node.description}\n{self.problem_text[:1000]}",
                        max_iterations=1,
                        score_threshold=8.0,
                        min_chars=1500,
                        checklist=self._get_rigor_checklist(),
                    )
                else:
                    print(f"    对公式进行Critique-Improvement...")
                    formulas = self.critique_engine.critique_and_improve(
                        content=formulas,
                        content_type="modeling",
                        context=f"任务: {task_node.description}\n{self.problem_text[:1000]}",
                        max_iterations=1,
                        score_threshold=8.0,
                        min_chars=1500,
                    )

            # 保存结果
            self.coordinator.save_task_result(
                task_id, {"analysis": task_analysis, "formulas": formulas}, key="modeling"
            )
            all_formulas.append(formulas)
            all_models.append(task_analysis)

        modeling_result = {
            "formulas": "\n\n".join(all_formulas),
            "models": "\n\n".join(all_models),
        }

        self.context["formulas"] = modeling_result["formulas"]
        self._save_json("stage_2_modeling/modeling.json", modeling_result)
        self._save_text("stage_2_modeling/formulas.md", modeling_result["formulas"])

        return modeling_result

    # ========================================================================
    # Stage 3: 计算求解
    # ========================================================================

    def _stage_computational_solving(self, modeling: Dict) -> Dict[str, Any]:
        """
        计算求解阶段

        明确使用 Claude CLI 生成代码
        使用 CodeExecutor 执行与调试
        """
        print("  [Stage 3] 开始计算求解...")

        formulas = modeling.get("formulas", "")

        # 设计算法（默认使用 Claude CLI）
        print("  [Stage 3.1] 设计求解算法...")
        algorithm_desc = self._design_algorithm_with_claude_cli(modeling)
        self.context["algorithm"] = algorithm_desc
        print(f"  [Stage 3.1] 算法设计完成 ({len(algorithm_desc)} 字符)")

        # 构建代码生成 Prompt
        # 注：要求代码精简，避免输出被截断
        code_prompt = f"""请编写精简的Python求解代码，解决以下数学建模问题。

【数学模型核心】
{formulas[:1500]}

【算法设计】
{algorithm_desc[:1000]}

【赛题摘要】
{self.problem_text[:1000]}

【数据文件】
{json.dumps(self.data_files, ensure_ascii=False, indent=2)}

硬性要求：
1. 代码必须精简（不超过200行），不要定义不必要的类，用函数即可
2. 代码开头必须是 import 语句，不要有任何中文说明文字
3. 代码必须能够读取上述数据文件（使用 pandas 读取 Excel）
4. 代码运行后必须使用以下代码写入结果：
   import os, json
   output_dir = os.environ.get('OUTPUT_DIR', '.')
   out_path = os.path.join(output_dir, 'execution', 'results.json')
   os.makedirs(os.path.dirname(out_path), exist_ok=True)
   with open(out_path, 'w') as f:
       json.dump(results, f, indent=2)
5. 代码中必须包含 main() 函数，且 if __name__ == '__main__': main()
6. 所有数值结果使用 Python 原生 float/int，不要嵌套numpy类型
7. 不得使用 input() 等交互式函数
8. 结果必须包含具体的数值，不能是空值或占位符
9. 代码必须在60秒内运行完成。禁止使用大规模网格搜索或复杂全局优化。如需搜索，网格点总数不得超过200个。
10. 优先使用解析解或简化公式，避免复杂数值积分和迭代优化。

输出纯Python代码，不要包含 markdown 代码块标记，不要任何解释。"""

        # 使用 CodeExecutor 生成代码
        # 注：代码生成使用 API（输出更稳定、长度限制宽松），代码修复使用 Claude CLI
        print("  [Stage 3.2] 生成求解代码...")
        result = self.code_executor.generate_and_run(
            prompt=code_prompt,
            system_prompt="你是Python编程专家，只输出代码，不输出任何解释。",
            data_files=self.data_files,
            filename="solve.py",
            use_claude_cli=False,
        )
        print(f"  [Stage 3.2] 代码生成完成，success={result.get('success', False)}")

        code = result.get("code", "")
        execution_result = result.get("execution_result", {})

        # 结果解读
        print("  解读计算结果...")
        interpretation = self._interpret_results(execution_result, modeling)

        # 保存
        self._save_text("stage_4_coding/solve.py", code)
        self._save_json("stage_5_execution/execution_result.json", execution_result)
        self._save_text("stage_6_result_analysis/interpretation.md", interpretation)

        return {
            "code": code,
            "execution_result": execution_result,
            "interpretation": interpretation,
            "success": result.get("success", False),
        }

    # ========================================================================
    # Stage 3b: 样本外回测/验证（CCF-A 强制）
    # ========================================================================
    #
    # 审稿意见核心拒稿点之一："No out-of-sample validation or backtesting is
    # provided... no quantitative scoring (RMSE/CRPS) of the scenario engine is
    # reported." 本阶段在 Stage 3 求解后，强制生成并运行一个回测脚本：
    # 用前期数据估计参数、对后期做点预测+区间预测，输出 RMSE/MAE/CRPS，
    # 结果写入 memory_pool["validation_summary"] 供论文 experiments 章引用。
    # 非 CCF-A 模板跳过本阶段。

    def _stage_validation_backtest(self, solving: Dict) -> Dict[str, Any]:
        """样本外回测阶段：生成并运行回测脚本，产出 RMSE/MAE/CRPS。"""
        if not self._is_ccf_a_template():
            return {"skipped": True, "reason": "non-CCF-A template"}

        print("  [Stage 3b] CCF-A 样本外回测/验证...")
        code = solving.get("code", "")
        execution_result = solving.get("execution_result", {})
        formulas = self.context.get("formulas", "")

        backtest_prompt = f"""请基于以下已运行的模型代码与结果，编写一个**样本外回测脚本** backtest.py。

【已运行模型代码（参考其数据加载与模型结构）】
{code[:2500]}

【模型核心公式】
{formulas[:1500]}

【已有计算结果（含可用数据键名）】
{json.dumps(execution_result, ensure_ascii=False, indent=2)[:1500]}

【数据文件】
{json.dumps(self.data_files, ensure_ascii=False, indent=2)}

硬性要求：
1. 必须切分 train/test（如时间序列用前期训练、后期验证；若无明确时间，用 K-fold）
2. 在 train 上估计参数，在 test 上做点预测 + 区间预测
3. 必须计算并输出 RMSE / MAE / CRPS（连续等级概率分数）三个指标
4. 若模型本身是情景模拟而非预测，则做"历史回填"：用历史观测校验模型对历史期的复现精度，并报告复现误差
5. 结果写入：
   import os, json
   output_dir = os.environ.get('OUTPUT_DIR', '.')
   out_path = os.path.join(output_dir, 'execution', 'backtest_metrics.json')
   os.makedirs(os.path.dirname(out_path), exist_ok=True)
   with open(out_path, 'w') as f:
       json.dump({{"rmse": ..., "mae": ..., "crps": ..., "train_period": ..., "test_period": ..., "per_metric": {{}}}}, f, indent=2)
6. 代码精简（<=150 行），开头是 import，含 main()，60 秒内完成
7. 若数据不足以做严格回测，输出 metrics 全为 None 并在 backtest_notes 字段说明原因

输出纯 Python 代码，不要 markdown 代码块标记。"""

        result = self.code_executor.generate_and_run(
            prompt=backtest_prompt,
            system_prompt="你是Python回测专家，只输出代码，不输出任何解释。",
            data_files=self.data_files,
            filename="backtest.py",
            use_claude_cli=False,
        )

        backtest_metrics = result.get("execution_result", {})
        backtest_code = result.get("code", "")

        # 也尝试从 execution/backtest_metrics.json 读取（脚本约定的写入路径）
        bt_path = self.output_dir / "execution" / "backtest_metrics.json"
        if bt_path.exists():
            try:
                with open(bt_path, "r", encoding="utf-8") as f:
                    backtest_metrics = json.load(f)
            except Exception:
                pass

        self._save_text("stage_5_execution/backtest.py", backtest_code)
        self._save_json("stage_5_execution/backtest_metrics.json", backtest_metrics)

        # 生成结构化摘要供论文 experiments 章引用
        summary = self._summarize_validation(backtest_metrics)
        if hasattr(self, "memory_pool") and isinstance(self.memory_pool, dict):
            self.memory_pool["validation_summary"] = summary
        self._save_text("stage_5_execution/validation_summary.md", summary)
        if self.memory_system:
            self.memory_system.store_shared("validation_summary", summary, metadata={"stage": "3b"})

        return {
            "backtest_code": backtest_code,
            "backtest_metrics": backtest_metrics,
            "validation_summary": summary,
            "success": result.get("success", False),
        }

    def _summarize_validation(self, metrics: Dict) -> str:
        """把回测指标提炼成论文可直接引用的摘要。"""
        if not metrics:
            return "（回测未产出指标，请在论文中诚实说明数据不足以做严格样本外验证。）"
        mtext = json.dumps(metrics, ensure_ascii=False, indent=2)[:1200]
        prompt = f"""请对以下样本外回测指标进行结构化提炼，生成一份"回测验证摘要"（300-400字），供论文 Experiments 章节直接引用。

【回测指标】
{mtext}

摘要要求：
1. 训练/测试期划分（1句话）
2. 三个核心指标数值：RMSE / MAE / CRPS（若有）
3. 各分项指标的对比（如有多目标，列出每个目标的误差）
4. 一句话结论：模型预测技能是否可接受（CRPS<0.3 或 RMSE<历史波动率一般可接受）
5. 若指标为 None，说明数据不足的原因与拟采取的替代验证（历史回填/敏感性）

输出纯文本，不要JSON。"""
        try:
            return _call_llm(prompt, "你是计量经济学与预测验证专家。")
        except Exception:
            return f"回测指标：{mtext}"

    # ========================================================================
    # Stage 3c: 复现性打包（CCF-A 强制）
    # ========================================================================
    #
    # 审稿意见核心拒稿点："code is not available for reproducibility...
    # Monte Carlo priors and parameter distributions are not fully documented
    # (sources, ranges, correlations)." 本阶段聚合代码、参数先验、协方差、
    # 随机种子、依赖快照、构造量对账表，落盘成 reproducibility_bundle/，
    # 并生成 memory_pool["reproducibility_summary"] 供 appendix 引用。

    def _stage_reproducibility_packaging(self, solving: Dict, validation: Dict) -> Dict[str, Any]:
        """复现性打包：代码/参数/种子/依赖/对账表聚合落盘。"""
        if not self._is_ccf_a_template():
            return {"skipped": True, "reason": "non-CCF-A template"}

        print("  [Stage 3c] CCF-A 复现性打包...")
        bundle_dir = self.output_dir / "reproducibility_bundle"
        bundle_dir.mkdir(parents=True, exist_ok=True)

        code = solving.get("code", "")
        execution_result = solving.get("execution_result", {})
        backtest_code = validation.get("backtest_code", "") if validation else ""
        formulas = self.context.get("formulas", "")

        # 1. 复制求解代码与回测代码
        (bundle_dir / "solve.py").write_text(code, encoding="utf-8")
        if backtest_code:
            (bundle_dir / "backtest.py").write_text(backtest_code, encoding="utf-8")

        # 2. 让 LLM 从模型公式与代码中抽取参数先验表 + 构造量对账表 + 随机种子
        repro_prompt = f"""请从以下模型公式与代码中抽取复现性所需的全部信息，输出严格 JSON。

【模型公式】
{formulas[:2500]}

【求解代码】
{code[:2500]}

【已有计算结果（含数据键名）】
{json.dumps(execution_result, ensure_ascii=False, indent=2)[:1500]}

请输出 JSON，包含以下字段：
{{
  "parameter_priors": [
    {{"name": "参数名", "symbol": "符号", "distribution": "分布族如 Normal/Uniform/LogNormal", "mean_or_range": "均值或范围", "source": "来源(文献/标定/假设)", "correlated_with": ["其他参数名"]}}
  ],
  "random_seeds": {{"main_seed": 42, "monte_carlo_seed": 123, "backtest_seed": 456}},
  "constructed_quantities": [
    {{"name": "如 土地净收入", "formula": "LaTeX公式", "raw_inputs": ["原始输入1","原始输入2"], "reconciliation": "从原始输入到使用值的对账步骤"}}
  ],
  "data_files": [{{"name": "文件名", "role": "训练/测试/校准", "preprocessing": "预处理步骤"}}],
  "compute_environment": {{"python": "3.x", "key_packages": ["numpy","pandas","..."], "hardware": "CPU/GPU"}}
}}

只输出 JSON，不要任何其他文字。"""

        try:
            repro_json_text = _call_llm(repro_prompt, "你是机器学习复现性专家，只输出JSON。")
            import re as _re
            m = _re.search(r'\{.*\}', repro_json_text, _re.DOTALL)
            repro_data = json.loads(m.group()) if m else {}
        except Exception as e:
            print(f"  [Stage 3c] 参数抽取失败: {e}")
            repro_data = {"error": str(e)}

        # 3. 固定随机种子声明
        seeds = repro_data.get("random_seeds", {"main_seed": 42, "monte_carlo_seed": 123, "backtest_seed": 456})
        (bundle_dir / "seeds.json").write_text(json.dumps(seeds, ensure_ascii=False, indent=2), encoding="utf-8")

        # 4. 参数先验表
        (bundle_dir / "parameter_priors.json").write_text(
            json.dumps(repro_data.get("parameter_priors", []), ensure_ascii=False, indent=2), encoding="utf-8")

        # 5. 构造量对账表
        (bundle_dir / "constructed_quantities.json").write_text(
            json.dumps(repro_data.get("constructed_quantities", []), ensure_ascii=False, indent=2), encoding="utf-8")

        # 6. 依赖快照
        try:
            import subprocess as _sp
            req = _sp.run(["pip", "freeze"], capture_output=True, text=True, timeout=30)
            (bundle_dir / "requirements_frozen.txt").write_text(req.stdout, encoding="utf-8")
        except Exception:
            (bundle_dir / "requirements_frozen.txt").write_text("# pip freeze 失败，请手动记录依赖", encoding="utf-8")

        # 7. README
        readme = f"""# Reproducibility Bundle (CCF-A)

本目录由系统自动生成，对应审稿意见的复现性要求。

## 文件清单
- `solve.py` — 主求解代码
- `backtest.py` — 样本外回测代码（若生成）
- `seeds.json` — 随机种子声明
- `parameter_priors.json` — 参数先验分布/范围/来源/相关性
- `constructed_quantities.json` — 构造量对账表（如"土地净收入"从原始输入到使用值）
- `requirements_frozen.txt` — 依赖快照

## 随机种子
{json.dumps(seeds, ensure_ascii=False, indent=2)}

## 复现步骤
1. 创建环境：`python -m venv venv && source venv/bin/activate && pip install -r requirements_frozen.txt`
2. 设置 `OUTPUT_DIR` 环境变量指向输出目录
3. 运行 `python solve.py`，结果写入 `$OUTPUT_DIR/execution/results.json`
4. 运行 `python backtest.py`（若有），回测指标写入 `$OUTPUT_DIR/execution/backtest_metrics.json`
"""
        (bundle_dir / "README.md").write_text(readme, encoding="utf-8")

        # 8. 生成结构化摘要供论文 appendix 引用
        summary = self._summarize_reproducibility(repro_data, bundle_dir)
        if hasattr(self, "memory_pool") and isinstance(self.memory_pool, dict):
            self.memory_pool["reproducibility_summary"] = summary
        self._save_text("reproducibility_bundle/reproducibility_summary.md", summary)
        if self.memory_system:
            self.memory_system.store_shared("reproducibility_summary", summary, metadata={"stage": "3c"})

        return {
            "bundle_dir": str(bundle_dir),
            "repro_data": repro_data,
            "reproducibility_summary": summary,
        }

    def _summarize_reproducibility(self, repro_data: Dict, bundle_dir: Path) -> str:
        """复现性打包摘要，供论文 appendix 直接引用。"""
        n_params = len(repro_data.get("parameter_priors", []))
        n_quant = len(repro_data.get("constructed_quantities", []))
        seeds = repro_data.get("random_seeds", {})
        seeds_str = ", ".join(f"{k}={v}" for k, v in seeds.items()) if seeds else "未声明"
        return (
            f"复现性打包已完成，产物位于 `{bundle_dir.name}/`。\n"
            f"- 参数先验表：{n_params} 个参数（分布/范围/来源/相关性）→ `parameter_priors.json`\n"
            f"- 构造量对账表：{n_quant} 个构造量（如土地净收入、政策乘数）→ `constructed_quantities.json`\n"
            f"- 随机种子：{seeds_str} → `seeds.json`\n"
            f"- 依赖快照 → `requirements_frozen.txt`\n"
            f"- 主代码 → `solve.py`，回测代码 → `backtest.py`\n"
            f"- 复现步骤见 `README.md`。"
        )

    def _design_algorithm(self, modeling: Dict) -> str:
        """设计求解算法"""
        formulas = modeling.get("formulas", "")

        prompt = f"""基于以下数学模型，设计详细的求解算法。

【数学模型】
{formulas[:2000]}

要求：
1. 选择合适的数值方法或优化算法
2. 给出详细的算法步骤（伪代码或分步说明）
3. 分析算法的时间复杂度和空间复杂度
4. 讨论算法的收敛性和稳定性
5. 设置关键参数及其选择依据

输出至少800字的算法设计文档。"""

        print("    调用LLM设计算法...")
        return _call_llm(prompt, "你是算法设计专家。")

    def _design_algorithm_with_claude_cli(self, modeling: Dict) -> str:
        """设计求解算法（默认走 API Provider，即 _call_llm）"""
        formulas = modeling.get("formulas", "")

        prompt = f"""基于以下数学模型，设计详细的求解算法。

【数学模型核心】
{formulas[:2000]}

【赛题摘要】
{self.problem_text[:1000]}

要求：
1. 选择合适的数值方法或优化算法
2. 给出详细的算法步骤（伪代码或分步说明）
3. 分析算法的时间复杂度和空间复杂度
4. 讨论算法的收敛性和稳定性
5. 设置关键参数及其选择依据

输出至少800字的算法设计文档。输出纯文本，不要包含 markdown 代码块标记。"""

        # 默认走 API Provider（_call_llm 已支持 API Provider，仅在无可用 API 时回退 CLI）
        print("    [Algorithm] 调用 LLM 设计算法...")
        try:
            return _call_llm(prompt, "你是算法设计专家。")
        except Exception as e:
            print(f"    [Algorithm] API 设计失败: {e}，回退到 Claude CLI...")
            if not self._claude_path:
                raise

            full_prompt = f"你是算法设计专家，擅长为数学建模问题设计高效求解算法。\n\n{prompt}"
            cmd = [
                self._claude_path,
                "-p",
                "--model", "sonnet",
                "--output-format", "json",
                full_prompt,
            ]
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
            )
            stdout, stderr = proc.communicate(timeout=300)

            if proc.returncode != 0:
                error_msg = stderr.decode("utf-8", errors="replace")
                raise RuntimeError(f"Claude CLI 算法设计失败: {error_msg[:500]}")

            stdout_text = stdout.decode("utf-8", errors="replace").strip()
            try:
                data = json.loads(stdout_text)
                result = data.get("result", "")
                if isinstance(result, str):
                    text = result.strip()
                    if text.startswith("```"):
                        lines = text.splitlines()
                        if lines[0].startswith("```"):
                            lines = lines[1:]
                        if lines and lines[-1].strip() == "```":
                            lines = lines[:-1]
                        text = "\n".join(lines).strip()
                    return text
            except json.JSONDecodeError:
                return stdout_text.strip()
            return stdout_text.strip()

    def _interpret_results(self, execution_result: Dict, modeling: Dict) -> str:
        """解读计算结果"""
        result_text = json.dumps(execution_result, ensure_ascii=False, indent=2)
        print(f"    计算结果大小: {len(result_text)} 字符")

        prompt = f"""请对以下计算结果进行专业解读。

【计算结果】
{result_text[:3000]}

【数学模型】
{modeling.get('formulas', '')[:1500]}

要求：
1. 总结主要数值结果
2. 分析结果的合理性和物理意义
3. 进行误差分析（如有可能）
4. 讨论结果的稳健性
5. 给出明确的研究结论

输出至少1500字的详细解读。"""

        return _call_llm(prompt, "你是数据分析专家。")

    # ========================================================================
    # Stage 4: 论文生成
    # ========================================================================

    def _stage_paper_generation_v2(self) -> str:
        print("  [Stage 4] 开始论文生成（分段逐章 + 记忆衔接）...")
        charts = self._generate_charts()
        self.context["charts"] = charts

        # Collect generated chart files for auto-insertion
        charts_dir = self.output_dir / "stage_7_charts"
        chart_files = [str(f) for f in sorted(charts_dir.glob("fig_*.png"))] if charts_dir.exists() else []
        if chart_files:
            print(f"  [Stage 4] 发现 {len(chart_files)} 张图表，将在论文中自动引用")

        memory_for_paper = {
            "analysis_summary": self.memory_pool.get("analysis_summary", ""),
            "modeling_summary": self.memory_pool.get("modeling_summary", ""),
            "algorithm_summary": self.memory_pool.get("algorithm_summary", ""),
            "results_summary": self.memory_pool.get("results_summary", ""),
            "validation_summary": self.memory_pool.get("validation_summary", ""),
            "reproducibility_summary": self.memory_pool.get("reproducibility_summary", ""),
        }
        paper = self.paper_generator.generate_paper_v2(
            context=self.context,
            memory_pool=memory_for_paper,
            use_critique=self.use_critique,
            chart_files=chart_files,
        )
        self.memory_pool["chapter_summaries"] = self.paper_generator.chapter_summaries
        self._save_json("final/chapter_summaries.json", self.paper_generator.chapter_summaries)
        paper_path = self.paper_generator.save_paper(paper, "MathModeling_Paper.md")
        print(f"  [Stage 4.3] 论文已保存: {paper_path}")
        return paper

    def _summarize_analysis(self, analysis: Dict) -> str:
        """生成阶段1的结构化摘要（供阶段2/4使用）"""
        analysis_text = json.dumps(analysis, ensure_ascii=False, indent=2)
        prompt = f"""请对以下问题分析结果进行结构化提炼，生成一份"分析摘要"（严格控制在400-500字）。

【原始分析】
{analysis_text[:2000]}

【摘要要求】
1. 问题背景（1句话）
2. 关键子问题（每题1句话，只列目标）
3. 核心假设（3条）
4. 解决思路（2-3句话）
5. 关键数值约束（列出3个最重要的）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数学建模专家，擅长提炼问题要点。")

    def _summarize_modeling(self, modeling: Dict) -> str:
        """生成阶段2的结构化摘要（供阶段3/4使用）"""
        formulas = modeling.get("formulas", "")[:1500]
        prompt = f"""请对以下数学建模结果进行结构化提炼，生成一份"建模摘要"（严格控制在400-500字）。

【原始建模内容】
{formulas}

【摘要要求】
1. 核心变量（列出5个最重要的变量、符号、单位）
2. 核心公式（1-2个最关键的LaTeX公式）
3. 求解策略（1句话）
4. 与子问题的映射（每子问题1句话）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数学建模专家，擅长提炼数学模型要点。")

    def _summarize_algorithm(self, solving: Dict) -> str:
        """生成阶段3算法摘要"""
        code = solving.get("code", "")[:1000]
        prompt = f"""请对以下求解算法进行结构化提炼（严格控制在200-300字）。

【代码片段】
{code}

【摘要要求】
1. 算法名称
2. 输入数据
3. 核心步骤（3步）
4. 输出格式

输出纯文本。"""
        return _call_llm(prompt, "你是算法专家。")

    def _summarize_results(self, solving: Dict) -> str:
        """生成阶段3结果摘要（供论文阶段使用）"""
        result = solving.get("execution_result", {})
        result_text = json.dumps(result, ensure_ascii=False, indent=2)[:1500]
        interpretation = solving.get("interpretation", "")[:1000]

        prompt = f"""请对以下计算结果进行结构化提炼，生成一份"结果摘要"（严格控制在400-500字）。

【计算结果】
{result_text}

【结果解读】
{interpretation}

【摘要要求】
1. 关键数值结果（用表格列出最重要的5个数值）
2. 主要发现（2-3句话）
3. 与模型的对应（1句话）
4. 可视化建议（1句话）

输出纯文本，不要JSON。"""
        return _call_llm(prompt, "你是数据分析专家，擅长提炼数值结果。")

    def _generate_charts(self) -> str:
        """
        生成论文图表

        Pipeline:
        1. Try ChartPipeline (LLM-driven with validation and retry)
        2. Fall back to TemplateChartGenerator (always produces output)
        """
        charts_dir = self.output_dir / "stage_7_charts"
        charts_dir.mkdir(parents=True, exist_ok=True)

        # Load results
        results_path = self.output_dir / "execution" / "results.json"
        R = {}
        if results_path.exists():
            with open(results_path, "r", encoding="utf-8") as f:
                R = json.load(f)

        if not R:
            print("  [Charts] results.json 为空，跳过图表生成")
            return f"图表保存于: {charts_dir}"

        # Step 1: Try LLM-driven ChartPipeline
        chart_count = 0
        try:
            from src.charts import ChartPipeline
            pipeline = ChartPipeline(self.output_dir, R, call_llm=_call_llm)
            chart_count = pipeline.run()
            if chart_count > 0:
                print(f"  [ChartPipeline] 成功生成 {chart_count} 张图表")
        except Exception as e:
            print(f"  [ChartPipeline] 失败: {e}")

        # Step 2: Fallback to template-based charts if pipeline produced nothing
        if chart_count == 0:
            print("  [Charts] ChartPipeline 未产出，使用模板图表生成...")
            try:
                from src.charts import TemplateChartGenerator
                gen = TemplateChartGenerator(self.output_dir, R)
                chart_count = gen.generate_all()
                if chart_count > 0:
                    print(f"  [TemplateCharts] 生成 {chart_count} 张模板图表")
            except Exception as e:
                print(f"  [TemplateCharts] 失败: {e}")

        if chart_count == 0:
            print("  [Charts] 未能生成任何图表，results.json 可能缺少可可视化数据")
        else:
            # List generated files
            png_files = list(charts_dir.glob("*.png"))
            print(f"  [Charts] 已生成 {len(png_files)} 张图表:")
            for f in sorted(png_files)[:6]:
                print(f"    - {f.name}")

        return f"图表保存于: {charts_dir} (共 {chart_count} 张)"


    # ========================================================================
    # 工具方法
    # ========================================================================

    def _convert_to_docx(self, md_path: Path):
        """将Markdown论文转换为docx格式"""
        docx_path = md_path.parent / "数学建模论文.docx"

        # 方法1: 使用 pandoc（质量最好）
        try:
            import subprocess
            result = subprocess.run(
                ["pandoc", str(md_path), "-o", str(docx_path), "--resource-path", str(md_path.parent)],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and docx_path.exists():
                print(f"  论文已导出为Word: {docx_path}")
                return
        except FileNotFoundError:
            pass
        except Exception as e:
            print(f"  pandoc导出失败: {e}")

        # 方法2: 使用 python-docx
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn

            doc = Document()

            # 设置中文字体支持
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 读取markdown内容
            text = md_path.read_text(encoding="utf-8")

            # 简单的markdown到docx转换
            lines = text.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]

                # 标题
                if line.startswith("# "):
                    p = doc.add_heading(line[2:], level=0)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith("#### "):
                    doc.add_heading(line[5:], level=3)
                # 表格（简单处理）
                elif line.startswith("|"):
                    # 收集表格行
                    table_lines = []
                    while i < len(lines) and lines[i].startswith("|"):
                        table_lines.append(lines[i])
                        i += 1
                    i -= 1  # 回退一行，外层循环会+1

                    if len(table_lines) >= 2:
                        # 解析表头
                        header_cells = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                        # 创建表格
                        table = doc.add_table(rows=1, cols=len(header_cells))
                        table.style = 'Light Grid Accent 1'
                        hdr_cells = table.rows[0].cells
                        for j, cell_text in enumerate(header_cells):
                            hdr_cells[j].text = cell_text
                        # 添加数据行（跳过第二行的分隔符）
                        for row_line in table_lines[2:]:
                            cells = [c.strip() for c in row_line.split("|")[1:-1]]
                            if len(cells) == len(header_cells):
                                row_cells = table.add_row().cells
                                for j, cell_text in enumerate(cells):
                                    row_cells[j].text = cell_text
                # 图片
                elif line.startswith("!["):
                    # 尝试提取图片路径
                    match = re.search(r'!\[.*?\]\((.*?)\)', line)
                    if match:
                        img_path = md_path.parent / match.group(1)
                        if img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5.5))
                            except Exception:
                                pass
                # 普通段落
                elif line.strip():
                    doc.add_paragraph(line)

                i += 1

            doc.save(str(docx_path))
            print(f"  论文已导出为Word: {docx_path}")
        except ImportError:
            print("  未安装 python-docx，跳过Word导出。可运行: pip install python-docx")
        except Exception as e:
            print(f"  Word导出失败: {e}")

    def _convert_to_tex(self, md_path: Path):
        # CCF-A 等模板：用 template_name 作为 template_id，让 tex_exporter 从
        # paper_templates 注册表取正确的 documentclass/cls_file/sty_files。
        template_dir = Path(__file__).parent.parent / "config" / "latex_templates" / "mcm"
        tex_path = md_path.parent / "MathModeling_Paper.tex"
        try:
            converter = MarkdownToTexConverter(
                template_dir, md_path.parent, template_id=self.template_name
            )
            converter.export(md_path, tex_path)
            print(f"  论文已导出为LaTeX: {tex_path}")
        except Exception as e:
            print(f"  LaTeX导出失败: {e}")

    def _save_json(self, rel_path: str, data: Any):
        path = self.output_dir / rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    def _save_text(self, rel_path: str, text: str):
        """保存文本文件"""
        path = self.output_dir / rel_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(text, encoding="utf-8")

    def _parse_json_safely(self, text: str, default: Any) -> Any:
        """安全解析JSON"""
        try:
            # 尝试提取 JSON 块
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                return json.loads(match.group())
            return json.loads(text)
        except Exception:
            return default

    def _default_analysis(self) -> Dict:
        """默认分析问题结果"""
        return {
            "background": "数学建模问题分析",
            "sub_problems": [
                {
                    "id": "task_1",
                    "description": "建立模型求解",
                    "objective": "解答问题",
                    "key_constraints": [],
                    "suggested_methods": ["数学建模方法"],
                }
            ],
            "key_assumptions": ["合理假设"],
            "data_dependencies": {},
            "solution_approach": "建立数学模型进行求解",
        }


def run_auto_paper_generation(
    problem_file: str = "problem.md",
    data_files: Dict[str, str] = None,
    output_dir: str = "work",
    template_name: str = "math_modeling",
) -> str:
    """运行全自动论文生成"""
    if data_files is None:
        data_files = {}

    problem_text = ""
    if Path(problem_file).exists():
        with open(problem_file, "r", encoding="utf-8") as f:
            problem_text = f.read()

    engine = UnifiedWorkflow(
        output_dir=output_dir,
        template_name=template_name,
    )

    return engine.run_full_workflow(
        problem_text=problem_text,
        data_files=data_files,
    )
